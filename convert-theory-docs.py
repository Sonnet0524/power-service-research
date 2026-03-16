#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
批量将理论文档Markdown转换为Word格式
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import os
from pathlib import Path

class MarkdownToWordConverter:
    def __init__(self):
        self.doc = None
        self.current_table = None
        
    def create_document(self):
        """创建新文档并设置默认字体"""
        self.doc = Document()
        
        # 设置默认中文字体
        style = self.doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        return self.doc
    
    def parse_markdown(self, content):
        """解析Markdown内容"""
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # 跳过Frontmatter
            if line.strip() == '---':
                i += 1
                while i < len(lines) and lines[i].strip() != '---':
                    i += 1
                i += 1
                continue
            
            # 标题 (H1)
            if line.startswith('# '):
                self.add_heading(line[2:], 1)
                i += 1
                continue
            
            # 标题 (H2)
            if line.startswith('## '):
                self.add_heading(line[3:], 2)
                i += 1
                continue
            
            # 标题 (H3)
            if line.startswith('### '):
                self.add_heading(line[4:], 3)
                i += 1
                continue
            
            # 标题 (H4)
            if line.startswith('#### '):
                self.add_heading(line[5:], 4)
                i += 1
                continue
            
            # 代码块
            if line.strip().startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                self.add_code_block('\n'.join(code_lines))
                i += 1
                continue
            
            # 表格
            if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
                table_lines = [line]
                i += 1
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                self.add_table(table_lines)
                continue
            
            # 列表
            if line.strip().startswith('- ') or line.strip().startswith('* '):
                self.add_list_item(line.strip()[2:], 0)
                i += 1
                continue
            
            if re.match(r'^\s+[-\*]\s', line):
                indent = len(line) - len(line.lstrip())
                level = indent // 2
                content = re.sub(r'^\s+[-\*]\s', '', line)
                self.add_list_item(content, level)
                i += 1
                continue
            
            # 空行
            if line.strip() == '':
                i += 1
                continue
            
            # 普通文本
            self.add_paragraph(line.strip())
            i += 1
    
    def add_heading(self, text, level):
        """添加标题"""
        if not text.strip():
            return
            
        heading = self.doc.add_heading(text, level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 设置字体
        for run in heading.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            if level == 1:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(27, 54, 93)  # 深蓝
                run.font.bold = True
            elif level == 2:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(27, 54, 93)
                run.font.bold = True
            elif level == 3:
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(46, 90, 76)  # 森林绿
                run.font.bold = True
            else:
                run.font.size = Pt(11)
                run.font.bold = True
    
    def add_paragraph(self, text):
        """添加段落"""
        if not text.strip():
            return
            
        # 处理Markdown加粗
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'__(.+?)__', r'\1', text)
        
        # 处理Markdown斜体
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'_(.+?)_', r'\1', text)
        
        p = self.doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(10.5)
    
    def add_list_item(self, text, level):
        """添加列表项"""
        # 处理加粗
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        
        indent = '    ' * level
        prefix = '● ' if level == 0 else '○ ' if level == 1 else '■ '
        
        p = self.doc.add_paragraph(f"{indent}{prefix}{text}")
        p.paragraph_format.left_indent = Inches(0.3 * (level + 1))
        
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(10.5)
    
    def add_code_block(self, code):
        """添加代码块/ASCII图"""
        if not code.strip():
            return
            
        p = self.doc.add_paragraph()
        run = p.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
        
        # 添加浅色背景效果
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
    
    def add_table(self, lines):
        """添加表格"""
        if len(lines) < 2:
            return
        
        # 解析表头
        headers = [cell.strip() for cell in lines[0].split('|') if cell.strip()]
        
        # 解析数据行
        data_rows = []
        for line in lines[2:]:  # 跳过表头和分隔线
            if '|' in line:
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                if cells:
                    data_rows.append(cells)
        
        if not data_rows:
            return
        
        # 创建表格
        table = self.doc.add_table(rows=len(data_rows)+1, cols=len(headers))
        table.style = 'Table Grid'
        
        # 设置表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    run.font.size = Pt(10)
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 设置数据行
        for row_idx, row_data in enumerate(data_rows):
            row_cells = table.rows[row_idx+1].cells
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(row_cells):
                    row_cells[col_idx].text = cell_text
                    for paragraph in row_cells[col_idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Microsoft YaHei'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                            run.font.size = Pt(9.5)
                    row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 添加表格后空行
        self.doc.add_paragraph()
    
    def convert_file(self, md_file_path, output_path):
        """转换单个Markdown文件"""
        print(f"正在转换: {md_file_path.name}")
        
        # 读取Markdown文件
        with open(md_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 创建文档
        self.create_document()
        
        # 解析并转换
        self.parse_markdown(content)
        
        # 保存
        self.doc.save(output_path)
        print(f"  ✓ 已生成: {output_path.name}")
        
        # 统计字数
        total_chars = sum(len(p.text) for p in self.doc.paragraphs)
        print(f"  字数统计: {total_chars} 字符")
        
        return total_chars


def main():
    """主函数"""
    # 路径配置
    theory_dir = Path('research/theory')
    output_dir = Path('output/theory-docs')
    
    # 确保输出目录存在
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # 获取所有Markdown文件
    md_files = sorted(theory_dir.glob('*.md'))
    
    print(f"发现 {len(md_files)} 个理论文档")
    print("="*50)
    
    # 文档名称映射
    name_mapping = {
        '2026-03-16-stationed-service-system-design.md': '01-驻点工作体系设计.docx',
        '2026-03-16-information-pre-positioning-implementation.md': '02-信息前置实施方案.docx',
        '2026-03-16-community-grid-integration-plan.md': '03-社网共建融合方案.docx',
        '2026-03-16-knowledge-base-construction-plan.md': '04-知识库建设方案.docx',
        '2026-03-16-stationed-workflow-design.md': '05-驻点工作流程设计.docx',
        '2026-03-16-information-usage-scenarios.md': '06-信息使用场景设计.docx',
        '2026-03-16-supplementary-content.md': '07-补充内容汇编.docx',
        '2026-03-16-digital-empowerment-theory.md': '08-数字化赋能理论.docx',
        '2026-03-16-enterprise-wechat-solution.md': '09-企业微信解决方案.docx',
        '2026-03-16-mobile-service-model.md': '10-移动服务模式研究.docx',
        '2026-03-16-completeness-analysis-framework.md': '11-完备性分析框架.docx',
        '2026-03-16-information-value-analysis.md': '12-信息价值分析.docx',
    }
    
    # 转换所有文件
    converter = MarkdownToWordConverter()
    total_chars_all = 0
    
    for md_file in md_files:
        if md_file.name in name_mapping:
            output_name = name_mapping[md_file.name]
        else:
            # 如果没有映射，使用原名
            output_name = md_file.stem + '.docx'
        
        output_path = output_dir / output_name
        
        try:
            chars = converter.convert_file(md_file, output_path)
            total_chars_all += chars
            print()
        except Exception as e:
            print(f"  ✗ 转换失败: {e}\n")
    
    print("="*50)
    print(f"✓ 全部转换完成!")
    print(f"✓ 输出目录: {output_dir}")
    print(f"✓ 文档总数: {len(md_files)} 个")
    print(f"✓ 总字数: {total_chars_all} 字符")


if __name__ == '__main__':
    main()
