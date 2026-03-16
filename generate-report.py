#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
武侯供电中心驻点工作研究报告Word生成器（增强版）
修复图形问题，增加视觉效果
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, Rectangle, Circle, FancyArrowPatch, ArrowStyle
import matplotlib.patches as mpatches
import numpy as np

# 设置中文字体
plt.rcParams['font.sans-serif'] = ['Arial Unicode MS', 'SimHei', 'Microsoft YaHei', 'Heiti TC', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False


class DiagramGenerator:
    """架构图生成器 - 修复版"""
    
    def __init__(self, output_dir):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def create_workflow_diagram(self):
        """创建驻点工作流程图 - 垂直布局"""
        fig, ax = plt.subplots(1, 1, figsize=(12, 14))
        ax.set_xlim(0, 12)
        ax.set_ylim(0, 14)
        ax.axis('off')
        
        # 标题
        ax.text(6, 13.5, '武侯供电中心驻点工作流程', fontsize=20, fontweight='bold',
                ha='center', va='center', color='#1a1a1a')
        
        # 定义三个阶段的参数（从上到下）
        phases = [
            {
                'name': '驻点前阶段',
                'color': '#1a5276',
                'y': 10.5,
                'time': 'T-1天',
                'items': [
                    '查询小区信息',
                    '获取重点客户清单',
                    '查询历史记录',
                    '制定工作计划',
                    '准备物资资料'
                ]
            },
            {
                'name': '驻点中阶段',
                'color': '#7b241c',
                'y': 6.5,
                'time': 'T日',
                'items': [
                    '公示牌张贴/检查',
                    '加入/检查微信群',
                    '重点客户走访',
                    '配电房检查',
                    '应急资源采集',
                    '其他工作'
                ]
            },
            {
                'name': '驻点后阶段',
                'color': '#1e8449',
                'y': 2.5,
                'time': 'T+1~T+7天',
                'items': [
                    '整理工作记录',
                    '跟踪问题整改',
                    '回访重点客户',
                    '处理客户诉求',
                    '更新知识库',
                    '总结归档'
                ]
            }
        ]
        
        # 绘制三个阶段
        for i, phase in enumerate(phases):
            # 主框
            rect_height = 2.8
            rect = FancyBboxPatch((1.5, phase['y']-rect_height/2), 9, rect_height,
                                   boxstyle="round,pad=0.05,rounding_size=0.2",
                                   facecolor=phase['color'], edgecolor='white', linewidth=3,
                                   alpha=0.95)
            ax.add_patch(rect)
            
            # 时间标签（左侧）
            ax.text(0.8, phase['y'], phase['time'], fontsize=11, fontweight='bold',
                   ha='center', va='center', color='#333333',
                   bbox=dict(boxstyle='round,pad=0.3', facecolor='#f0f0f0', edgecolor='#999999'))
            
            # 阶段名称
            ax.text(6, phase['y']+0.9, phase['name'], fontsize=14, fontweight='bold',
                   ha='center', va='center', color='white')
            
            # 工作内容（使用项目符号）
            content_text = '  •  ' + '\n  •  '.join(phase['items'])
            ax.text(6, phase['y']-0.3, content_text, fontsize=9, ha='center', va='center',
                   color='white', alpha=0.95, linespacing=1.6)
            
            # 绘制向下箭头（除了最后一个）
            if i < len(phases) - 1:
                next_y = phases[i+1]['y'] + rect_height/2 + 0.3
                current_y = phase['y'] - rect_height/2 - 0.3
                
                # 绘制粗箭头
                ax.annotate('', xy=(6, next_y), xytext=(6, current_y),
                           arrowprops=dict(arrowstyle='->', color='#333333', lw=3,
                                          connectionstyle="arc3,rad=0"))
        
        plt.tight_layout()
        filepath = os.path.join(self.output_dir, 'workflow.png')
        plt.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white', pad_inches=0.3)
        plt.close()
        return filepath
    
    def create_three_domain_diagram(self):
        """创建三大领域协同模型图 - 水平布局"""
        fig, ax = plt.subplots(1, 1, figsize=(14, 9))
        ax.set_xlim(0, 14)
        ax.set_ylim(0, 9)
        ax.axis('off')
        
        # 标题
        ax.text(7, 8.5, '三大领域协同模型', fontsize=20, fontweight='bold',
                ha='center', va='center', color='#1a1a1a')
        
        # 三个领域的数据
        domains = [
            {
                'name': '网格服务',
                'subtitle': '服务前置',
                'color': '#1a5276',
                'x': 2.5,
                'desc': '让客户找到人\n压降95598工单',
                'arrow_text': '客户信息',
                'arrow_color': '#1a5276'
            },
            {
                'name': '用电检查',
                'subtitle': '信息前置',
                'color': '#7b241c',
                'x': 7,
                'desc': '掌握设备状况\n消除安全隐患',
                'arrow_text': '设备信息',
                'arrow_color': '#7b241c'
            },
            {
                'name': '应急保供',
                'subtitle': '资源前置',
                'color': '#1e8449',
                'x': 11.5,
                'desc': '建立知识库\n快速应急响应',
                'arrow_text': '',
                'arrow_color': '#1e8449'
            }
        ]
        
        # 绘制三个领域框
        box_width = 3.5
        box_height = 3.0
        
        for i, domain in enumerate(domains):
            y_center = 5.0
            
            # 主框
            rect = FancyBboxPatch((domain['x']-box_width/2, y_center-box_height/2), 
                                   box_width, box_height,
                                   boxstyle="round,pad=0.05,rounding_size=0.2",
                                   facecolor=domain['color'], edgecolor='white', linewidth=3,
                                   alpha=0.95)
            ax.add_patch(rect)
            
            # 领域名称
            ax.text(domain['x'], y_center+0.8, domain['name'], fontsize=14, fontweight='bold',
                   ha='center', va='center', color='white')
            
            # 副标题（黄色高亮）
            ax.text(domain['x'], y_center+0.2, domain['subtitle'], fontsize=12, fontweight='bold',
                   ha='center', va='center', color='#FFD700')
            
            # 描述
            ax.text(domain['x'], y_center-0.8, domain['desc'], fontsize=9, ha='center', va='center',
                   color='white', alpha=0.95, linespacing=1.5)
            
            # 绘制向右箭头（除了最后一个）
            if i < len(domains) - 1:
                next_x = domains[i+1]['x'] - box_width/2 - 0.2
                current_x = domain['x'] + box_width/2 + 0.2
                arrow_y = y_center + 0.2
                
                # 绘制箭头
                ax.annotate('', xy=(next_x, arrow_y), xytext=(current_x, arrow_y),
                           arrowprops=dict(arrowstyle='->', color=domain['arrow_color'], lw=2.5))
                
                # 箭头标签
                mid_x = (current_x + next_x) / 2
                ax.text(mid_x, arrow_y+0.3, domain['arrow_text'], fontsize=9, 
                       ha='center', va='bottom', color=domain['arrow_color'], fontweight='bold')
        
        # 绘制返回箭头（从应急保供到网格服务的弧线）
        # 起点：应急保供左下角
        start_x = domains[2]['x'] - box_width/2
        start_y = 5.0 - box_height/2 - 0.2
        # 终点：网格服务右下角
        end_x = domains[0]['x'] + box_width/2
        end_y = 5.0 - box_height/2 - 0.2
        
        # 使用弧线连接
        ax.annotate('', xy=(end_x+0.1, end_y), xytext=(start_x-0.1, start_y),
                   arrowprops=dict(arrowstyle='->', color='#1e8449', lw=2.5,
                                  connectionstyle="arc3,rad=-0.4"))
        
        # 弧线标签
        ax.text(7, 1.3, '提升信心', fontsize=10, ha='center', va='center', 
               color='#1e8449', fontweight='bold',
               bbox=dict(boxstyle='round,pad=0.2', facecolor='white', edgecolor='#1e8449', alpha=0.9))
        
        # 底部核心价值
        ax.text(7, 0.5, '驻点工作核心价值：服务前置 + 信息前置 + 资源前置', 
               fontsize=12, fontweight='bold', ha='center', va='center', color='#333333',
               bbox=dict(boxstyle='round,pad=0.3', facecolor='#fff9e6', edgecolor='#d4ac0d'))
        
        plt.tight_layout()
        filepath = os.path.join(self.output_dir, 'three_domain.png')
        plt.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white', pad_inches=0.3)
        plt.close()
        return filepath
    
    def create_information_flow_diagram(self):
        """创建信息前置机制流程图 - 水平线性布局"""
        fig, ax = plt.subplots(1, 1, figsize=(14, 7))
        ax.set_xlim(0, 14)
        ax.set_ylim(0, 7)
        ax.axis('off')
        
        # 标题
        ax.text(7, 6.5, '信息前置机制流程', fontsize=20, fontweight='bold',
                ha='center', va='center', color='#1a1a1a')
        
        # 流程步骤数据
        steps = [
            {'name': '数据源', 'color': '#5d6d7e', 'x': 2, 
             'desc': '营销系统\nPMS系统\n95598系统'},
            {'name': '数据汇聚', 'color': '#1a5276', 'x': 5,
             'desc': '小区台账\n客户台账\n工作台账'},
            {'name': '信息推送', 'color': '#7b241c', 'x': 8,
             'desc': '重点客户清单\n自动推送'},
            {'name': '现场使用', 'color': '#1e8449', 'x': 11,
             'desc': '快速查询\n精准服务'},
            {'name': '数据更新', 'color': '#b7950b', 'x': 8, 'y': 2,
             'desc': '一次录入\n自动联动', 'is_bottom': True}
        ]
        
        box_width = 2.2
        box_height = 1.8
        
        # 绘制主要流程框（上方）
        for i in range(4):
            step = steps[i]
            y = 4.0
            
            # 主框
            rect = FancyBboxPatch((step['x']-box_width/2, y-box_height/2), 
                                   box_width, box_height,
                                   boxstyle="round,pad=0.05,rounding_size=0.15",
                                   facecolor=step['color'], edgecolor='white', linewidth=2.5,
                                   alpha=0.95)
            ax.add_patch(rect)
            
            # 步骤名称
            ax.text(step['x'], y+0.5, step['name'], fontsize=11, fontweight='bold',
                   ha='center', va='center', color='white')
            
            # 描述
            ax.text(step['x'], y-0.3, step['desc'], fontsize=8, ha='center', va='center',
                   color='white', alpha=0.95, linespacing=1.4)
            
            # 绘制向右箭头
            if i < 3:
                next_step = steps[i+1]
                start_x = step['x'] + box_width/2 + 0.1
                end_x = next_step['x'] - box_width/2 - 0.1
                
                ax.annotate('', xy=(end_x, y), xytext=(start_x, y),
                           arrowprops=dict(arrowstyle='->', color='#666666', lw=2))
        
        # 绘制下方返回框
        bottom_step = steps[4]
        bottom_y = 2.0
        
        rect = FancyBboxPatch((bottom_step['x']-box_width/2, bottom_y-box_height/2), 
                               box_width, box_height,
                               boxstyle="round,pad=0.05,rounding_size=0.15",
                               facecolor=bottom_step['color'], edgecolor='white', linewidth=2.5,
                               alpha=0.95)
        ax.add_patch(rect)
        
        ax.text(bottom_step['x'], bottom_y+0.5, bottom_step['name'], fontsize=11, fontweight='bold',
               ha='center', va='center', color='white')
        ax.text(bottom_step['x'], bottom_y-0.3, bottom_step['desc'], fontsize=8, ha='center', va='center',
               color='white', alpha=0.95, linespacing=1.4)
        
        # 绘制从"现场使用"到"数据更新"的向下箭头
        ax.annotate('', xy=(bottom_step['x']+box_width/2+0.3, bottom_y), 
                   xytext=(steps[3]['x'], 4.0-box_height/2-0.2),
                   arrowprops=dict(arrowstyle='->', color='#1e8449', lw=2,
                                  connectionstyle="arc3,rad=0.2"))
        
        # 绘制从"数据更新"到"数据源"的返回箭头（弧线）
        ax.annotate('', xy=(steps[0]['x'], 4.0-box_height/2-0.3), 
                   xytext=(bottom_step['x']-box_width/2-0.2, bottom_y),
                   arrowprops=dict(arrowstyle='->', color='#b7950b', lw=2,
                                  connectionstyle="arc3,rad=0.3"))
        
        # 标签
        ax.text(3, 1.0, '数据回流', fontsize=9, ha='center', va='center', 
               color='#666666', style='italic',
               bbox=dict(boxstyle='round,pad=0.2', facecolor='#f5f5f5', edgecolor='#999999'))
        
        ax.text(12.5, 3.0, '闭环\n管理', fontsize=9, ha='center', va='center', 
               color='#1e8449', fontweight='bold',
               bbox=dict(boxstyle='round,pad=0.2', facecolor='#e8f8f5', edgecolor='#1e8449'))
        
        plt.tight_layout()
        filepath = os.path.join(self.output_dir, 'information_flow.png')
        plt.savefig(filepath, dpi=200, bbox_inches='tight', facecolor='white', pad_inches=0.3)
        plt.close()
        return filepath


class WordDocumentGenerator:
    """Word文档生成器 - 增强版"""
    
    def __init__(self):
        self.doc = None
        self.setup_styles()
    
    def setup_styles(self):
        """设置文档样式"""
        self.doc = Document()
        
        # 设置文档默认字体
        self.doc.styles['Normal'].font.name = '微软雅黑'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        self.doc.styles['Normal'].font.size = Pt(11)
        
    def add_title(self, text, level=0):
        """添加标题"""
        if level == 0:
            # 主标题
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        else:
            # 子标题
            heading = self.doc.add_heading(text, level=level)
            heading.style.font.name = '微软雅黑'
            heading.style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            if level == 1:
                heading.style.font.size = Pt(18)
                heading.style.font.color.rgb = RGBColor(0, 51, 102)
            elif level == 2:
                heading.style.font.size = Pt(16)
                heading.style.font.color.rgb = RGBColor(0, 76, 153)
            else:
                heading.style.font.size = Pt(14)
                heading.style.font.color.rgb = RGBColor(51, 102, 153)
    
    def add_paragraph(self, text, bold=False, italic=False, color=None):
        """添加段落"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
        return p
    
    def add_highlight_paragraph(self, text, keyword_list):
        """添加带高亮的段落"""
        p = self.doc.add_paragraph()
        
        # 分段处理，高亮关键词
        remaining = text
        for keyword in keyword_list:
            if keyword in remaining:
                parts = remaining.split(keyword, 1)
                if parts[0]:
                    run = p.add_run(parts[0])
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(11)
                
                # 高亮关键词
                run = p.add_run(keyword)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = RGBColor(192, 0, 0)  # 红色
                
                remaining = parts[1] if len(parts) > 1 else ""
        
        if remaining:
            run = p.add_run(remaining)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(11)
        
        return p
    
    def add_key_point(self, emoji, title, content):
        """添加关键点（带emoji）"""
        p = self.doc.add_paragraph()
        
        # Emoji
        run = p.add_run(f"{emoji} ")
        run.font.size = Pt(12)
        
        # 标题
        run = p.add_run(f"{title}：")
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 76, 153)
        
        # 内容
        run = p.add_run(content)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)
        
        return p
    
    def add_bullet_point(self, text, level=0):
        """添加项目符号"""
        p = self.doc.add_paragraph(text, style='List Bullet' if level == 0 else 'List Number')
        for run in p.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.size = Pt(11)
        return p
    
    def add_image(self, image_path, width=Inches(6)):
        """添加图片"""
        self.doc.add_picture(image_path, width=width)
        last_paragraph = self.doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()
    
    def save(self, filename):
        """保存文档"""
        self.doc.save(filename)


def generate_report():
    """生成完整的研究报告 - 增强版"""
    output_dir = '/Users/sonnet/opencode/power-service-research/output'
    
    # 生成图表
    print("正在生成图表...")
    diagram_gen = DiagramGenerator(output_dir)
    workflow_img = diagram_gen.create_workflow_diagram()
    three_domain_img = diagram_gen.create_three_domain_diagram()
    info_flow_img = diagram_gen.create_information_flow_diagram()
    print("✅ 图表生成完成")
    
    # 生成Word文档
    print("正在生成Word文档...")
    gen = WordDocumentGenerator()
    
    # ===== 封面 =====
    gen.add_title("🏢 武侯供电中心现场驻点工作研究报告")
    gen.add_paragraph("")
    gen.add_paragraph("")
    gen.add_paragraph("📋 研究课题：现场供电服务研究", bold=True)
    gen.add_paragraph("🎯 研究对象：2026年武侯供电中心现场驻点工作任务清单")
    gen.add_paragraph("📅 研究时间：2026年3月16日")
    gen.add_paragraph("🔬 研究方法：SEARCH-R方法论")
    gen.add_paragraph("📊 研究深度：Level 0-2（第一性原理到设计原则）")
    gen.add_paragraph("")
    gen.add_paragraph("")
    gen.add_paragraph("👥 研究团队：王丹、宋戈、Research Agent", bold=True)
    gen.add_paragraph("🔒 报告密级：内部")
    gen.add_paragraph("📌 报告版本：v1.0")
    
    gen.add_page_break()
    
    # ===== 执行摘要 =====
    gen.add_title("📋 执行摘要", level=1)
    
    gen.add_paragraph(
        "本报告基于SEARCH-R方法论，对武侯供电中心现场驻点工作进行了系统性研究。通过深入分析研究对象，"
        "借鉴40余个供电服务创新实践案例和10篇学术文献，构建了完整的驻点工作体系。研究表明，武侯供电中心的"
        "驻点工作应定位为", bold=False
    )
    
    p = gen.doc.add_paragraph()
    run = p.add_run("流动式驻点")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    
    run = p.add_run("，核心在于通过")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    
    run = p.add_run("信息前置")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    
    run = p.add_run("实现精准服务，最终目标是")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    
    run = p.add_run("压降95598工单")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    
    run = p.add_run("、提升客户满意度。")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    
    gen.add_paragraph(
        "研究明确了三大专业领域（网格服务、用电检查、应急保供）的协同关系，设计了完整的工作流程和实施方案，"
        "并提出了信息前置机制、社网共建模式、知识库建设等关键举措。本报告为武侯供电中心提升服务水平提供了"
        "系统性的理论支撑和可操作的实施方案。"
    )
    
    gen.add_page_break()
    
    # ===== 第一章 =====
    gen.add_title("第一章 研究背景与目标", level=1)
    
    gen.add_title("1.1 研究背景", level=2)
    gen.add_paragraph(
        "城区供电中心开展现场驻点工作，已形成工作任务清单，覆盖三大专业领域：网格服务、用电检查、应急保供。"
        "然而，随着电力服务要求的不断提升，传统的驻点工作模式面临新的挑战："
    )
    
    gen.add_key_point("❓", "核心问题", "如何在有限的时间内实现最大的服务价值？如何通过驻点工作有效压降95598工单？如何建立三大专业领域的协同机制？")
    
    gen.add_paragraph(
        "在此背景下，本研究旨在系统性地评估和完善驻点工作体系，明确工作定位，优化工作流程，"
        "建立协同机制，为武侯供电中心提升服务水平提供科学指导。"
    )
    
    gen.add_title("1.2 研究目标", level=2)
    gen.add_paragraph("本研究围绕四个核心问题展开：")
    
    gen.add_key_point("1️⃣", "完备性问题", "现场驻点工作是否梳理完备？需要从流动式驻点的视角重新审视工作清单，识别缺失项、冗余项和需调整项。")
    gen.add_key_point("2️⃣", "价值挖掘问题", "驻点工作能为电力保供和客户服务提供什么？需要从服务前置、信息前置、资源前置三个维度挖掘价值。")
    gen.add_key_point("3️⃣", "关系梳理问题", "三大专业领域之间的内在关系是什么？需要建立协同模型，明确数据流转机制。")
    gen.add_key_point("4️⃣", "提升路径问题", "如何通过现场驻点提升公司服务水平？需要设计系统性的提升方案和实施路径。")
    
    gen.add_page_break()
    
    # ===== 第二章 =====
    gen.add_title("第二章 研究方法与过程", level=1)
    
    gen.add_title("2.1 SEARCH-R方法论", level=2)
    
    p = gen.doc.add_paragraph()
    run = p.add_run("本研究采用")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    
    run = p.add_run("SEARCH-R方法论")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    run = p.add_run("，这是一个包含七个阶段的系统性研究循环：")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    
    gen.add_paragraph(
        "Survey（观察调研）→ Explore（探索检索）→ Analyze（分析思考）→ Review（评审探讨）→ "
        "Confirm（确认验证）→ Harvest（收获产出）→ Reflect（反思迭代）"
    )
    
    gen.add_paragraph(
        "在Survey阶段，研究团队深入理解了《2026年武侯供电中心现场驻点工作任务清单》和《营销稽查管控作业指导书》，"
        "建立了对研究对象的全面认知。在Explore阶段，系统检索了40余个供电服务创新实践案例和10篇学术文献，"
        "为研究提供了坚实的理论基础和实践参考。在Analyze阶段，构建了完整的理论框架和实施方案。"
    )
    
    gen.add_title("2.2 研究过程", level=2)
    gen.add_paragraph(
        "研究历时一天，经历了四个完整的会话周期。通过系统性的分析和设计，形成了完整的解决方案。"
    )
    
    gen.add_page_break()
    
    # ===== 第三章 =====
    gen.add_title("第三章 核心研究发现", level=1)
    
    gen.add_title("3.1 武侯模式的本质：流动式驻点", level=2)
    
    p = gen.doc.add_paragraph()
    run = p.add_run("通过对比分析其他地区供电服务实践，研究发现武侯供电中心的驻点工作应定位为")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    
    run = p.add_run("流动式驻点")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    
    run = p.add_run("。这与阿尔山天池景区、炎陵大院景区的驻点模式高度相似，核心特征在于时间有限、很久才去第二次、"
        "必须信息驱动、必须一次见效。"
    )
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    
    gen.add_paragraph(
        "流动式驻点不同于长期驻扎，不能依赖长期关系建立；不同于季节性驻点，不能依靠稳定的服务预期。"
        "其价值在于以信息前置实现精准服务，每次驻点都要有明确的成效。"
    )
    
    gen.add_title("3.2 四大核心原则", level=2)
    
    gen.add_key_point("⏱️", "时间有限原则", "要求筛选一次可完成的工作，单次驻点时长控制在2.5到4小时。")
    gen.add_key_point("📊", "信息驱动原则", "强调不能依赖关系驱动，必须通过信息前置实现精准服务。")
    gen.add_key_point("🎯", "抓住重点原则", "要求通过清单指导，精准聚焦高价值工作。")
    gen.add_key_point("✅", "一次见效原则", "要求每次驻点都要有明确的成效，建立跟踪反馈机制。")
    
    gen.add_page_break()
    
    # ===== 第四章 =====
    gen.add_title("第四章 驻点工作体系设计", level=1)
    
    gen.add_title("4.1 三大领域协同模型", level=2)
    
    p = gen.doc.add_paragraph()
    run = p.add_run("研究建立了")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    
    run = p.add_run("网格服务、用电检查、应急保供")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    run = p.add_run("三大领域的协同模型。")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    
    gen.add_paragraph(
        "网格服务聚焦服务前置，核心目标是压降95598工单，通过公示客户经理联系方式、加入小区微信群、走访重点客户等方式，"
        "让客户能够直接找到人，实现快速响应。"
    )
    
    gen.add_paragraph(
        "用电检查聚焦信息前置，核心目标是掌握设备状况，通过配电房检查、设备信息采集、隐患发现等方式，"
        "消除安全隐患，同时为应急保供提供数据支撑。"
    )
    
    gen.add_paragraph(
        "应急保供聚焦资源前置，核心目标是建立现场知识库，通过收集应急资源信息、记录接线图和接入点等方式，"
        "支撑快速故障定位和应急响应。"
    )
    
    gen.add_paragraph(
        "三大领域之间形成数据流转闭环：网格服务的客户信息支撑用电检查聚焦重点，"
        "用电检查的设备信息支撑应急保供知识库建设，应急保供的快速响应能力反过来提升网格服务的客户信心。"
    )
    
    gen.add_image(three_domain_img)
    
    gen.add_page_break()
    
    gen.add_title("4.2 工作清单设计", level=2)
    
    gen.add_paragraph(
        "基于流动式驻点的四大原则，研究设计了分级分类的工作清单。工作按照优先级分为三级："
    )
    
    gen.add_key_point("🔴", "极高优先级", "共8项工作，包括公示牌张贴、加入微信群、重点客户走访、配电房检查、应急资源信息采集等，预计耗时2.5小时，属于必做工作。")
    gen.add_key_point("🟡", "高优先级", "共4项工作，包括停电通知发布、电费催缴提醒、安全用电宣传、隐患整改跟踪，预计耗时45分钟，属于优先做工作。")
    gen.add_key_point("🟢", "中优先级", "共5项工作，包括客户诉求收集、能效建议提供等，预计耗时35分钟，属于有余力时做工作。")
    
    gen.add_paragraph(
        "工作清单的设计体现了时间有限原则，每项工作都有明确的时长控制；体现了抓住重点原则，"
        "极高优先级工作都是一次可完成、有明确产出的工作；体现了一次见效原则，"
        "每项工作都能直接产生服务价值或数据价值。"
    )
    
    gen.add_page_break()
    
    # ===== 第五章 =====
    gen.add_title("第五章 驻点工作完整流程", level=1)
    
    gen.add_title("5.1 流程概述", level=2)
    
    p = gen.doc.add_paragraph()
    run = p.add_run("完整的驻点工作流程分为")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    
    run = p.add_run("三个阶段")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    run = p.add_run("：驻点前阶段、驻点中阶段、驻点后阶段。驻点前阶段在T-1天完成，主要任务是信息准备和任务规划；驻点中阶段在T日完成，"
        "主要任务是现场执行和数据采集；驻点后阶段在T+1到T+7天完成，主要任务是跟踪回访和总结归档。"
    )
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    
    gen.add_image(workflow_img)
    
    gen.add_title("5.2 驻点前阶段", level=2)
    gen.add_paragraph(
        "驻点前阶段的核心是信息准备。客户经理需要查询小区基础信息，包括小区名称、地址、总户数、"
        "配电房位置等；获取系统自动推送的重点客户清单，包括敏感客户、频繁停电客户、特殊群体等；"
        "查询历史驻点记录，了解上次工作内容和遗留问题；制定详细的驻点工作计划，"
        "明确工作目标、走访计划、时间安排；准备所需物资资料。"
    )
    
    gen.add_title("5.3 驻点中阶段", level=2)
    gen.add_paragraph(
        "驻点中阶段的核心是现场执行。到达小区后，首先进行公示牌张贴或检查，确保客户经理联系方式对客户可见；"
        "然后加入或检查小区微信群，建立线上服务渠道；接着开展重点客户走访，根据客户类型提供差异化服务；"
        "随后进行配电房检查，发现设备隐患并记录；最后采集应急资源信息，完善知识库建设。"
    )
    
    gen.add_title("5.4 驻点后阶段", level=2)
    gen.add_paragraph(
        "驻点后阶段的核心是跟踪总结。需要整理工作记录，完善驻点工作台账，上传现场照片；"
        "跟踪问题整改，建立整改台账，确保发现的隐患得到及时处理；回访重点客户，了解诉求解决情况和满意度；"
        "处理客户诉求，建立闭环管理机制；更新知识库，将新采集的信息纳入知识库体系；最后进行总结归档。"
    )
    
    gen.add_page_break()
    
    # ===== 第六章 =====
    gen.add_title("第六章 关键机制设计", level=1)
    
    gen.add_title("6.1 信息前置机制", level=2)
    
    p = gen.doc.add_paragraph()
    run = p.add_run("信息前置机制是流动式驻点的")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    
    run = p.add_run("核心支撑")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    
    run = p.add_run("。该机制通过建立小区台账、客户台账、驻点工作台账三大数据表，"
        "实现信息的系统化管理。在驻点前1天，系统自动向客户经理推送重点客户清单，"
        "帮助客户经理提前了解服务对象，制定针对性的服务策略。"
    )
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    
    gen.add_paragraph(
        "在现场，客户经理可以通过企业微信或WPS多维表快速查询客户信息、配电房信息、应急资源信息，"
        "实现精准服务。服务完成后，通过一次录入机制，工作记录自动关联更新相关数据表。"
    )
    
    gen.add_image(info_flow_img)
    
    gen.add_title("6.2 社网共建模式", level=2)
    
    gen.add_paragraph(
        "社网共建是深化供电服务的重要途径。该模式通过组织融合、网格融合、服务融合三个层面，"
        "实现供电网格与社区网格的深度融合。在组织层面，建立供电所党支部与社区党组织的联建共建机制；"
        "在网格层面，实现供电网格与社区网格的范围对接、人员对接、信息共享；"
        "在服务层面，在社区党群服务中心设立电力服务驿站，提供基础办电服务和增值能效服务。"
    )
    
    gen.add_title("6.3 知识库建设", level=2)
    
    gen.add_paragraph(
        "应急保供知识库是快速响应的基础。知识库包括设备信息库、应急资源库、工作模板库三大组成部分。"
        "设备信息库记录配电房位置、设备参数、接线图等信息，支撑快速故障定位；"
        "应急资源库记录应急发电车接入点、抢修设备位置、应急人员联系方式等信息，支撑快速应急响应；"
        "工作模板库提供标准化的检查表、记录表、流程模板，支撑规范化作业。"
    )
    
    gen.add_page_break()
    
    # ===== 第七章 =====
    gen.add_title("第七章 实施路径与建议", level=1)
    
    gen.add_title("7.1 四阶段实施路径", level=2)
    
    gen.add_paragraph("建议采用四阶段实施路径：")
    
    gen.add_key_point("📅", "第一阶段（1-2个月）", "基础建设阶段，完成数据表设计、试点小区信息录入、信息推送机制建立等工作。")
    gen.add_key_point("🧪", "第二阶段（3-4个月）", "试点应用阶段，选择3到5个小区试点应用新机制，收集反馈并优化。")
    gen.add_key_point("🚀", "第三阶段（5-6个月）", "全面推广阶段，在所有小区全面应用新机制，建立考核机制。")
    gen.add_key_point("📈", "第四阶段（7-12个月）", "深化拓展阶段，推进社网共建，建设知识库，形成可复制推广的模式。")
    
    gen.add_title("7.2 关键成功因素", level=2)
    
    gen.add_paragraph("确保方案成功实施的关键在于五个方面：")
    
    gen.add_key_point("✓", "数据完整准确", "这是信息前置机制有效运行的基础。")
    gen.add_key_point("✓", "采集便捷高效", "提供便捷的工具和方法，降低客户经理的工作负担。")
    gen.add_key_point("✓", "应用广泛深入", "通过培训和考核确保方案的落地执行。")
    gen.add_key_point("✓", "持续更新维护", "定期更新数据，确保信息的时效性。")
    gen.add_key_point("✓", "安全可靠", "加强数据安全管理，保护客户隐私。")
    
    gen.add_title("7.3 预期成效", level=2)
    
    p = gen.doc.add_paragraph()
    run = p.add_run("通过实施本方案，预期将取得以下")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    
    run = p.add_run("显著成效")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 153, 0)
    
    run = p.add_run("：")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    
    gen.add_key_point("🎯", "服务成效", "实现居民办电不出社区，客户诉求响应时间缩短50%以上，客户满意度提升至95分以上，95598工单压降30%以上。")
    gen.add_key_point("⚙️", "管理成效", "实现资源整合，信息畅通，协同高效，考核科学。")
    gen.add_key_point("🏆", "品牌成效", "供电服务形象提升，形成社网共建服务品牌，建立可复制推广的经验。")
    
    gen.add_page_break()
    
    # ===== 第八章 =====
    gen.add_title("第八章 研究结论", level=1)
    
    gen.add_paragraph(
        "本研究通过系统性的分析和设计，为武侯供电中心现场驻点工作提供了完整的解决方案。"
        "研究明确了武侯模式属于流动式驻点，提出了四大核心原则，"
        "建立了三大领域协同模型，设计了完整的工作流程和实施方案。"
    )
    
    gen.add_paragraph(
        "研究的核心贡献在于将分散的驻点工作整合为系统化的工作体系，通过信息前置机制实现精准服务，"
        "通过社网共建模式深化服务触角，通过知识库建设支撑快速响应。这些机制相互配合，"
        "形成了从服务前置到信息前置再到资源前置的完整服务链条。"
    )
    
    p = gen.doc.add_paragraph()
    run = p.add_run("建议武侯供电中心按照四阶段实施路径推进方案落地，优先完善")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    
    run = p.add_run("信息前置机制")
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    
    run = p.add_run("，这是流动式驻点最关键的支撑；逐步建设知识库，先在重点小区试点；稳步推进社网共建，"
        "先建立联系机制再设立服务驿站；持续优化数字化工具，根据使用反馈不断改进。"
    )
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    
    gen.add_paragraph(
        "相信通过系统性的改革和创新，武侯供电中心的现场驻点工作将实现质的飞跃，"
        "为公司服务水平提升做出重要贡献，为其他地区供电服务提供可借鉴的经验。"
    )
    
    gen.add_page_break()
    
    # ===== 附录 =====
    gen.add_title("附录：文档清单", level=1)
    gen.add_paragraph("本研究报告包含以下配套文档：")
    
    docs = [
        ("一、《基于流动服务模式的驻点工作体系设计》", "理论框架文档，阐述核心定位、协同模型、工作清单等。"),
        ("二、《信息前置机制实施方案》", "详细的数据架构、推送机制、查询功能设计方案。"),
        ("三、《社网共建实施方案》", "社网共建模式、服务驿站设置、协同工作机制设计。"),
        ("四、《应急保供知识库建设方案》", "知识库架构、数据采集、应用场景设计。"),
        ("五、《驻点工作完整流程设计》", "T-1到T+7全流程详细步骤和检查清单。"),
        ("六、《信息获取与使用场景矩阵》", "信息分类、使用场景、优化建议。"),
        ("七、《驻点工作体系补充内容》", "工具指南、表单模板、培训方案、风险预案、考核细则、FAQ。")
    ]
    
    for title, desc in docs:
        p = gen.doc.add_paragraph()
        run = p.add_run(title)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 76, 153)
        
        run = p.add_run(f"：{desc}")
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(11)
    
    # 保存文档
    output_file = os.path.join(output_dir, '武侯供电中心驻点工作研究报告.docx')
    gen.save(output_file)
    
    print(f"\n✅ Word文档生成完成：{output_file}")
    print(f"✅ 配套图表已保存至：{output_dir}")
    print(f"📊 文档包含：8章内容 + 3张架构图 + 附录")
    print(f"🎨 增强功能：✓ 修复图形 ✓ 关键文字加粗 ✓ 颜色标记 ✓ Emoji图标")
    
    return output_file


if __name__ == '__main__':
    generate_report()
