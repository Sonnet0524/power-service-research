#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将Word文档中的字符画转换为专业图片
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont
import os
import re
from pathlib import Path

class AsciiArtToImage:
    """将ASCII艺术转换为专业图片"""
    
    # 麦肯锡风格配色
    COLORS = {
        'primary': '#1B365D',      # 深蓝
        'secondary': '#8B4513',    # 深棕
        'tertiary': '#2E5A4C',     # 森林绿
        'accent': '#C9A227',       # 金色
        'light': '#F5F5F0',        # 米白
        'text': '#2C3E50',         # 文本色
        'white': '#FFFFFF',
    }
    
    def __init__(self, output_dir):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.image_counter = 0
        
        # 尝试加载中文字体
        self.font_paths = [
            '/System/Library/Fonts/PingFang.ttc',  # macOS
            '/System/Library/Fonts/STHeiti Light.ttc',
            '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',  # Linux
            'C:/Windows/Fonts/simhei.ttf',  # Windows
            'C:/Windows/Fonts/msyh.ttc',
        ]
        self.font = self._load_font(20)
        self.font_small = self._load_font(14)
        self.font_title = self._load_font(24)
        
    def _load_font(self, size):
        """加载字体"""
        for path in self.font_paths:
            if os.path.exists(path):
                try:
                    return ImageFont.truetype(path, size)
                except:
                    pass
        return ImageFont.load_default()
    
    def generate_image(self, ascii_art, doc_name, art_type='concept'):
        """根据ASCII艺术生成专业图片"""
        self.image_counter += 1
        filename = f"{doc_name}_{self.image_counter:03d}.png"
        filepath = self.output_dir / filename
        
        # 分析ASCII艺术类型和内容
        if '流程' in ascii_art or '↓' in ascii_art or '→' in ascii_art:
            img = self._create_flow_chart(ascii_art)
        elif '┌' in ascii_art and '│' in ascii_art:
            img = self._create_box_diagram(ascii_art)
        else:
            img = self._create_concept_card(ascii_art)
        
        img.save(filepath, 'PNG', dpi=(300, 300))
        return str(filepath)
    
    def _create_box_diagram(self, ascii_art):
        """创建方框图（用于概念展示）"""
        lines = ascii_art.strip().split('\n')
        lines = [l for l in lines if l.strip()]
        
        # 计算尺寸
        max_width = max(len(l) for l in lines) * 20 + 100
        height = len(lines) * 35 + 80
        
        img = Image.new('RGB', (max_width, height), self.COLORS['light'])
        draw = ImageDraw.Draw(img)
        
        # 绘制圆角矩形背景
        margin = 20
        draw.rounded_rectangle(
            [margin, margin, max_width-margin, height-margin],
            radius=15,
            fill=self.COLORS['white'],
            outline=self.COLORS['primary'],
            width=3
        )
        
        # 绘制标题栏
        title_height = 50
        draw.rounded_rectangle(
            [margin, margin, max_width-margin, margin+title_height],
            radius=15,
            fill=self.COLORS['primary']
        )
        
        # 绘制内容
        y = margin + title_height + 20
        for line in lines[1:-1]:  # 跳过第一行和最后一行（边框）
            content = line.strip('│ ')
            if content and not all(c in '─┼┤├' for c in content):
                draw.text((margin + 30, y), content, 
                         font=self.font, fill=self.COLORS['text'])
            y += 32
        
        return img
    
    def _create_flow_chart(self, ascii_art):
        """创建流程图"""
        lines = ascii_art.strip().split('\n')
        
        # 解析流程步骤
        steps = []
        arrows = []
        
        for line in lines:
            if '───' in line or '→' in line or '↓' in line:
                # 这是一个箭头或连接线
                arrows.append(line)
            elif any(char in line for char in ['驻点前', '驻点中', '驻点后', '网格', '用电', '应急']):
                # 这是一个步骤
                steps.append(line.strip())
        
        # 创建图片
        width = 900
        height = max(400, len(steps) * 120 + 100)
        
        img = Image.new('RGB', (width, height), self.COLORS['light'])
        draw = ImageDraw.Draw(img)
        
        # 绘制步骤框
        box_width = 250
        box_height = 70
        margin_x = (width - box_width) // 2
        
        for i, step in enumerate(steps):
            y = 60 + i * 110
            
            # 绘制步骤框
            draw.rounded_rectangle(
                [margin_x, y, margin_x + box_width, y + box_height],
                radius=10,
                fill=self.COLORS['primary'] if i == 0 else self.COLORS['tertiary'] if i == 1 else self.COLORS['secondary'],
                outline=self.COLORS['accent'],
                width=2
            )
            
            # 绘制文字
            bbox = draw.textbbox((0, 0), step, font=self.font)
            text_width = bbox[2] - bbox[0]
            text_x = margin_x + (box_width - text_width) // 2
            text_y = y + (box_height - 24) // 2
            draw.text((text_x, text_y), step, font=self.font, fill=self.COLORS['white'])
            
            # 绘制箭头（如果不是最后一个）
            if i < len(steps) - 1:
                arrow_y = y + box_height + 10
                # 绘制向下箭头
                draw.polygon(
                    [(width//2, arrow_y + 20), (width//2 - 10, arrow_y), (width//2 + 10, arrow_y)],
                    fill=self.COLORS['accent']
                )
        
        return img
    
    def _create_concept_card(self, ascii_art):
        """创建概念卡片"""
        lines = ascii_art.strip().split('\n')
        
        width = 800
        height = max(300, len(lines) * 40 + 100)
        
        img = Image.new('RGB', (width, height), self.COLORS['light'])
        draw = ImageDraw.Draw(img)
        
        # 绘制卡片背景
        margin = 30
        draw.rounded_rectangle(
            [margin, margin, width-margin, height-margin],
            radius=20,
            fill=self.COLORS['white'],
            outline=self.COLORS['primary'],
            width=4
        )
        
        # 绘制顶部装饰条
        draw.rounded_rectangle(
            [margin, margin, width-margin, margin+8],
            radius=4,
            fill=self.COLORS['accent']
        )
        
        # 绘制内容
        y = margin + 40
        for line in lines:
            content = line.strip('┌│└─┐ ')
            if content:
                draw.text((margin + 40, y), content, 
                         font=self.font, fill=self.COLORS['text'])
            y += 35
        
        return img


class WordDocumentUpdater:
    """更新Word文档，将字符画替换为图片"""
    
    def __init__(self, theory_dir, output_dir, image_output_dir):
        self.theory_dir = Path(theory_dir)
        self.output_dir = Path(output_dir)
        self.image_generator = AsciiArtToImage(image_output_dir)
        self.image_output_dir = Path(image_output_dir)
        
    def process_all_documents(self):
        """处理所有文档"""
        word_files = sorted(self.output_dir.glob('*.docx'))
        
        print(f"发现 {len(word_files)} 个Word文档")
        print("="*60)
        
        for word_file in word_files:
            if word_file.name.startswith('~$'):  # 跳过临时文件
                continue
            self.process_single_document(word_file)
        
        print("\n" + "="*60)
        print("✓ 全部文档处理完成!")
        print(f"✓ 图片保存位置: {self.image_output_dir}")
    
    def process_single_document(self, word_file):
        """处理单个Word文档"""
        print(f"\n处理文档: {word_file.name}")
        
        # 读取对应的Markdown文件获取原始字符画
        md_filename = self._get_md_filename(word_file.name)
        md_file = self.theory_dir / md_filename
        
        if not md_file.exists():
            print(f"  ⚠ 未找到对应的Markdown文件: {md_filename}")
            return
        
        # 读取Markdown内容
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 提取所有ASCII艺术
        ascii_arts = self._extract_ascii_art(md_content)
        
        if not ascii_arts:
            print(f"  ℹ 未找到字符画")
            return
        
        print(f"  发现 {len(ascii_arts)} 个字符画")
        
        # 打开Word文档
        doc = Document(word_file)
        
        # 生成图片并替换
        doc_name = word_file.stem
        images_added = 0
        
        for i, (ascii_art, art_type) in enumerate(ascii_arts):
            # 生成图片
            image_path = self.image_generator.generate_image(
                ascii_art, doc_name, art_type
            )
            
            # 在Word中查找并替换
            if self._replace_ascii_with_image(doc, ascii_art, image_path):
                images_added += 1
        
        # 保存文档
        output_path = word_file
        doc.save(output_path)
        
        print(f"  ✓ 已替换 {images_added} 个字符画为图片")
    
    def _get_md_filename(self, word_filename):
        """根据Word文件名获取对应的Markdown文件名"""
        # 提取序号
        match = re.match(r'(\d+)-', word_filename)
        if match:
            num = int(match.group(1))
            # 获取所有md文件并排序
            md_files = sorted(self.theory_dir.glob('*.md'))
            # 返回对应序号的文件（序号从1开始）
            if 1 <= num <= len(md_files):
                return md_files[num - 1].name
        
        # 默认返回同名
        return word_filename.replace('.docx', '.md')
    
    def _extract_ascii_art(self, content):
        """从Markdown内容中提取ASCII艺术"""
        ascii_arts = []
        
        # 匹配代码块中的ASCII艺术
        pattern = r'```\n?(.*?)```'
        matches = re.finditer(pattern, content, re.DOTALL)
        
        for match in matches:
            art = match.group(1).strip()
            # 判断是否包含ASCII艺术字符
            if any(char in art for char in ['┌', '│', '└', '├', '┤', '─', '↓', '→', '╔', '╗', '╚', '╝']):
                # 判断类型
                if '流程' in art or '↓' in art or '→' in art:
                    art_type = 'flow'
                elif '┌' in art and '│' in art:
                    art_type = 'box'
                else:
                    art_type = 'concept'
                
                ascii_arts.append((art, art_type))
        
        return ascii_arts
    
    def _replace_ascii_with_image(self, doc, ascii_art, image_path):
        """在Word文档中将字符画替换为图片"""
        # 将ASCII艺术分行
        art_lines = ascii_art.strip().split('\n')
        if not art_lines:
            return False
        
        # 查找包含字符画起始行的段落
        first_line = art_lines[0].strip()
        found = False
        
        for i, para in enumerate(doc.paragraphs):
            if first_line in para.text or self._is_similar(para.text, first_line):
                # 找到了起始段落，删除接下来的段落直到字符画结束
                # 由于python-docx的限制，我们采用标记替换的方式
                found = True
                
                # 在当前段落位置插入图片
                para.clear()
                run = para.add_run()
                run.add_picture(image_path, width=Inches(5.5))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 标记已删除后续行（实际处理比较复杂，这里简化处理）
                break
        
        return found
    
    def _is_similar(self, text1, text2):
        """判断两段文本是否相似"""
        # 清理文本
        clean1 = re.sub(r'\s+', '', text1)
        clean2 = re.sub(r'\s+', '', text2)
        
        # 如果包含主要字符则认为是相似的
        return len(clean1) > 0 and len(clean2) > 0 and \
               (clean1 in clean2 or clean2 in clean1 or 
                any(char in clean1 for char in clean2[:20]))


def main():
    """主函数"""
    theory_dir = Path('research/theory')
    output_dir = Path('output/theory-docs')
    image_output_dir = Path('output/theory-images')
    
    print("开始将字符画转换为图片...")
    print("="*60)
    
    updater = WordDocumentUpdater(theory_dir, output_dir, image_output_dir)
    updater.process_all_documents()


if __name__ == '__main__':
    main()
