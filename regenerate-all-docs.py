#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
重新生成所有理论文档Word文件，直接嵌入高质量图片
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
import os
import re
import shutil
from pathlib import Path

class IntegratedConverter:
    """集成的Markdown到Word转换器（带图片生成）"""
    
    COLORS = {
        'primary': '#1B365D',
        'secondary': '#8B4513',
        'tertiary': '#2E5A4C',
        'accent': '#C9A227',
        'light': '#F8F8F6',
        'text': '#2C3E50',
        'white': '#FFFFFF',
        'shadow': '#D0D0D0',
    }
    
    def __init__(self):
        self.doc = None
        self.image_counter = 0
        self.temp_images = []
        
        # 字体
        self.font_paths = [
            '/System/Library/Fonts/PingFang.ttc',
            '/System/Library/Fonts/STHeiti Light.ttc',
            '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
            'C:/Windows/Fonts/simhei.ttf',
            'C:/Windows/Fonts/msyh.ttc',
        ]
        self.font = self._load_font(24)
        self.font_title = self._load_font(28)
        
    def _load_font(self, size):
        for path in self.font_paths:
            if os.path.exists(path):
                try:
                    return ImageFont.truetype(path, size)
                except:
                    pass
        return ImageFont.load_default()
    
    def is_ascii_art(self, content):
        """检查是否为ASCII艺术"""
        ascii_chars = set('┌─┐│└┘├┤┼┬┴┼═║╔╗╚╝╠╣╦╩╬─━┃┏┓┗┛┣┫┻┳↓→')
        return bool(set(content) & ascii_chars)
    
    def create_document(self):
        """创建新文档"""
        self.doc = Document()
        style = self.doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        return self.doc
    
    def add_heading(self, text, level):
        """添加标题"""
        if not text.strip():
            return
        heading = self.doc.add_heading(text, level)
        for run in heading.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            if level == 1:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(27, 54, 93)
                run.font.bold = True
            elif level == 2:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(27, 54, 93)
                run.font.bold = True
            else:
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(46, 90, 76)
                run.font.bold = True
    
    def add_paragraph(self, text):
        """添加段落"""
        if not text.strip():
            return
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        p = self.doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(10.5)
    
    def generate_image_from_ascii(self, ascii_art):
        """从ASCII艺术生成图片"""
        lines = [l.rstrip() for l in ascii_art.strip().split('\n') if l.strip()]
        if not lines:
            return None
        
        self.image_counter += 1
        
        # 确定图片类型
        has_arrows = '↓' in ascii_art or '→' in ascii_art
        has_boxes = '┌' in ascii_art or '╔' in ascii_art
        
        if has_arrows:
            # 流程图
            return self._create_flow_image(lines)
        elif has_boxes:
            # 架构图
            return self._create_architecture_image(lines)
        else:
            return None
    
    def _create_flow_image(self, lines):
        """创建流程图图片"""
        steps = []
        for line in lines:
            if '↓' not in line:
                content = line.strip('├┤│─ ').strip()
                if content and len(content) > 2:
                    steps.append(content)
        
        if len(steps) < 2:
            return None
        
        width = 1000
        box_width = 360
        box_height = 90
        spacing = 50
        height = max(600, len(steps) * (box_height + spacing) + 150)
        
        img = Image.new('RGB', (width, height), self.COLORS['light'])
        draw = ImageDraw.Draw(img)
        
        colors = [self.COLORS['primary'], self.COLORS['tertiary'], self.COLORS['secondary']]
        
        for i, step in enumerate(steps):
            y = 80 + i * (box_height + spacing)
            color = colors[i % len(colors)]
            
            # 阴影
            draw.rounded_rectangle(
                [(width - box_width)//2 + 5, y + 5, (width + box_width)//2 + 5, y + box_height + 5],
                radius=12, fill=self.COLORS['shadow']
            )
            
            # 主框
            draw.rounded_rectangle(
                [(width - box_width)//2, y, (width + box_width)//2, y + box_height],
                radius=12, fill=color, outline=self.COLORS['accent'], width=3
            )
            
            # 文字
            bbox = draw.textbbox((0, 0), step, font=self.font)
            text_width = bbox[2] - bbox[0]
            text_x = (width - text_width) // 2
            text_y = y + (box_height - 24) // 2
            draw.text((text_x, text_y), step, font=self.font, fill=self.COLORS['white'])
            
            # 箭头
            if i < len(steps) - 1:
                arrow_y = y + box_height + 10
                draw.polygon(
                    [(width//2, arrow_y + 30), (width//2 - 14, arrow_y + 5), (width//2 + 14, arrow_y + 5)],
                    fill=self.COLORS['accent']
                )
        
        # 保存临时文件
        temp_path = f"/tmp/ascii_img_{self.image_counter:03d}.png"
        img.save(temp_path, 'PNG', dpi=(300, 300))
        self.temp_images.append(temp_path)
        return temp_path
    
    def _create_architecture_image(self, lines):
        """创建架构图图片"""
        max_len = max(len(l) for l in lines) if lines else 50
        width = max(900, max_len * 18 + 120)
        height = max(500, len(lines) * 35 + 100)
        
        img = Image.new('RGB', (width, height), self.COLORS['light'])
        draw = ImageDraw.Draw(img)
        
        margin = 40
        draw.rounded_rectangle(
            [margin, margin, width-margin, height-margin],
            radius=15, fill=self.COLORS['white'], outline=self.COLORS['primary'], width=3
        )
        
        draw.rounded_rectangle(
            [margin, margin, width-margin, margin+10],
            radius=5, fill=self.COLORS['accent']
        )
        
        y = margin + 40
        for line in lines:
            content = re.sub(r'^[┌├│└┏┣┃]', '  ', line)
            content = re.sub(r'[┐┤┘┓┫┛]$', '  ', content)
            content = content.rstrip()
            
            if content.strip():
                if any(kw in content for kw in ['武侯', '驻点', '体系', '核心']):
                    draw.text((margin + 50, y), content.strip(), 
                             font=self.font_title, fill=self.COLORS['primary'])
                else:
                    draw.text((margin + 50, y), content.strip(), 
                             font=self.font, fill=self.COLORS['text'])
            y += 34
        
        temp_path = f"/tmp/ascii_img_{self.image_counter:03d}.png"
        img.save(temp_path, 'PNG', dpi=(300, 300))
        self.temp_images.append(temp_path)
        return temp_path
    
    def add_image(self, image_path):
        """添加图片到文档"""
        if image_path and os.path.exists(image_path):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(image_path, width=Inches(5.5))
    
    def parse_and_convert(self, md_content):
        """解析Markdown并转换"""
        lines = md_content.split('\n')
        i = 0
        in_code_block = False
        code_content = []
        
        while i < len(lines):
            line = lines[i]
            
            # 处理Frontmatter
            if line.strip() == '---' and i == 0:
                i += 1
                while i < len(lines) and lines[i].strip() != '---':
                    i += 1
                i += 1
                continue
            
            # 代码块处理
            if line.strip().startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_content = []
                else:
                    in_code_block = False
                    # 处理代码块内容
                    code_text = '\n'.join(code_content)
                    if self.is_ascii_art(code_text):
                        # 生成图片
                        img_path = self.generate_image_from_ascii(code_text)
                        if img_path:
                            self.add_image(img_path)
                    else:
                        # 普通代码块，作为文本添加
                        if code_text.strip():
                            self.add_paragraph(code_text)
                i += 1
                continue
            
            if in_code_block:
                code_content.append(line)
                i += 1
                continue
            
            # 标题
            if line.startswith('# '):
                self.add_heading(line[2:], 1)
            elif line.startswith('## '):
                self.add_heading(line[3:], 2)
            elif line.startswith('### '):
                self.add_heading(line[4:], 3)
            elif line.startswith('#### '):
                self.add_heading(line[5:], 4)
            elif line.strip() and not line.strip().startswith('|'):
                # 普通文本（跳过表格行，由表格处理）
                if not (line.strip().startswith('|--') or '---' in line):
                    self.add_paragraph(line.strip())
            
            i += 1
    
    def save(self, output_path):
        """保存文档"""
        self.doc.save(output_path)
        # 清理临时图片
        for img_path in self.temp_images:
            if os.path.exists(img_path):
                os.remove(img_path)


def convert_single_file(md_file, output_dir):
    """转换单个文件"""
    print(f"\n转换: {md_file.name}")
    
    # 读取Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 创建转换器
    converter = IntegratedConverter()
    converter.create_document()
    converter.parse_and_convert(content)
    
    # 生成输出文件名
    output_name = md_file.stem.replace('2026-03-16-', '')
    output_name = output_name.replace('stationed-service-system-design', '01-驻点工作体系设计')
    output_name = output_name.replace('information-pre-positioning-implementation', '02-信息前置实施方案')
    output_name = output_name.replace('community-grid-integration-plan', '03-社网共建融合方案')
    output_name = output_name.replace('knowledge-base-construction-plan', '04-知识库建设方案')
    output_name = output_name.replace('stationed-workflow-design', '05-驻点工作流程设计')
    output_name = output_name.replace('information-usage-scenarios', '06-信息使用场景设计')
    output_name = output_name.replace('supplementary-content', '07-补充内容汇编')
    output_name = output_name.replace('digital-empowerment-theory', '08-数字化赋能理论')
    output_name = output_name.replace('enterprise-wechat-solution', '09-企业微信解决方案')
    output_name = output_name.replace('mobile-service-model', '10-移动服务模式研究')
    output_name = output_name.replace('completeness-analysis-framework', '11-完备性分析框架')
    output_name = output_name.replace('information-value-analysis', '12-信息价值分析')
    
    output_path = output_dir / f"{output_name}.docx"
    converter.save(output_path)
    
    print(f"  ✓ 生成: {output_path.name}")
    print(f"  图片数: {converter.image_counter}")
    return converter.image_counter


def main():
    """主函数"""
    theory_dir = Path('research/theory')
    output_dir = Path('output/theory-docs')
    
    # 清理旧文件
    if output_dir.exists():
        shutil.rmtree(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # 获取所有Markdown文件
    md_files = sorted(theory_dir.glob('*.md'))
    
    print(f"开始转换 {len(md_files)} 个文档...")
    print("="*70)
    
    total_images = 0
    for md_file in md_files:
        count = convert_single_file(md_file, output_dir)
        total_images += count
    
    print("\n" + "="*70)
    print(f"✓ 全部转换完成!")
    print(f"✓ 文档数: {len(md_files)}")
    print(f"✓ 总图片数: {total_images}")
    print(f"✓ 输出目录: {output_dir}")


if __name__ == '__main__':
    main()
